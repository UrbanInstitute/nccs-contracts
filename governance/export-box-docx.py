#!/usr/bin/env python3
"""Render governance markdown to committee-ready Word docs for Box.

Usage:  python3 governance/export-box-docx.py <input.md> <output.docx>
Requires: python-docx. Git is the source of truth; these exports are
the committee-facing surface (see governance/README.md).
"""
import re, sys, pathlib
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def build(md_path, out_path):
    src = pathlib.Path(md_path).read_text()
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'; st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15
    for name, size, color, before in [('Heading 1', 16, '1F3864', 6),
                                      ('Heading 2', 12.5, '1F3864', 14),
                                      ('Heading 3', 11, '2E5496', 10)]:
        h = doc.styles[name]
        h.font.name = 'Calibri'; h.font.size = Pt(size); h.font.bold = True
        h.font.color.rgb = RGBColor.from_string(color)
        h.paragraph_format.space_before = Pt(before)
        h.paragraph_format.space_after = Pt(4)
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.8)
        sec.left_margin = sec.right_margin = Inches(0.9)

    def add_runs(p, text):
        token = re.compile(r'(\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*'
                           r'|(?<!\*)\*([^*\n]+)\*(?!\*)|`([^`]+)`)')
        pos = 0
        for m in token.finditer(text):
            if m.start() > pos:
                p.add_run(text[pos:m.start()])
            if m.group(2):
                rid = p.part.relate_to(
                    m.group(3),
                    docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                    is_external=True)
                hl = OxmlElement('w:hyperlink'); hl.set(qn('r:id'), rid)
                r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
                u = OxmlElement('w:u'); u.set(qn('w:val'), 'single')
                rPr.append(u)
                c = OxmlElement('w:color'); c.set(qn('w:val'), '0563C1')
                rPr.append(c)
                r.append(rPr)
                t2 = OxmlElement('w:t'); t2.text = m.group(2)
                t2.set(qn('xml:space'), 'preserve')
                r.append(t2); hl.append(r); p._p.append(hl)
            elif m.group(4):
                r = p.add_run(m.group(4)); r.bold = True
            elif m.group(5):
                r = p.add_run(m.group(5)); r.italic = True
            elif m.group(6):
                r = p.add_run(m.group(6))
                r.font.name = 'Consolas'; r.font.size = Pt(9.5)
            pos = m.end()
        if pos < len(text):
            p.add_run(text[pos:])

    def bg(cell, hexc):
        sh = OxmlElement('w:shd')
        sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexc)
        cell._tc.get_or_add_tcPr().append(sh)

    lines = src.split('\n'); i = 0
    while i < len(lines):
        ln = lines[i]
        if (ln.startswith('|') and i + 1 < len(lines)
                and set(lines[i+1].replace('|', '').strip()) <= set('-: ')):
            hdr = [c.strip() for c in ln.strip('|').split('|')]
            rows = []; i += 2
            while i < len(lines) and lines[i].startswith('|'):
                rows.append([c.strip() for c in lines[i].strip('|').split('|')])
                i += 1
            tbl = doc.add_table(rows=1 + len(rows), cols=len(hdr))
            tbl.style = 'Table Grid'
            for j, c in enumerate(hdr):
                cell = tbl.rows[0].cells[j]; p = cell.paragraphs[0]
                add_runs(p, c)
                for r in p.runs:
                    r.bold = True; r.font.size = Pt(9.5)
                bg(cell, 'D9E2F3')
            for ri, row in enumerate(rows):
                for j, c in enumerate(row):
                    p = tbl.rows[ri+1].cells[j].paragraphs[0]
                    add_runs(p, c)
                    for r in p.runs:
                        r.font.size = Pt(9.5)
            doc.add_paragraph(); continue
        m = re.match(r'^(#{1,3}) (.*)', ln)
        if m:
            doc.add_heading('', level=len(m.group(1)))
            add_runs(doc.paragraphs[-1], m.group(2))
        elif re.match(r'^- \[[ x]\] ', ln) or re.match(r'^[-*] ', ln):
            item = re.sub(r'^- \[([ x])\] ',
                          lambda m: ('☑ ' if m.group(1) == 'x'
                                     else '☐ '), ln)
            item = re.sub(r'^[-*] ', '', item)
            while (i + 1 < len(lines) and lines[i+1].startswith('  ')
                   and not re.match(r'^\s*([-*]|\d+\.|- \[)', lines[i+1])):
                i += 1; item += ' ' + lines[i].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, item); p.paragraph_format.space_after = Pt(3)
        elif re.match(r'^\d+\. ', ln):
            item = re.sub(r'^\d+\. ', '', ln)
            while (i + 1 < len(lines) and lines[i+1].startswith('   ')
                   and not re.match(r'^\s*(\d+\. |[-*] )', lines[i+1])):
                i += 1; item += ' ' + lines[i].strip()
            p = doc.add_paragraph(style='List Number')
            add_runs(p, item); p.paragraph_format.space_after = Pt(3)
        elif ln.strip() == '---':
            pass
        elif ln.strip():
            para = ln
            while (i + 1 < len(lines) and lines[i+1].strip()
                   and not re.match(r'^(#|\||[-*] |\d+\. |---|- \[)',
                                    lines[i+1])):
                i += 1; para += ' ' + lines[i].strip()
            p = doc.add_paragraph(); add_runs(p, para)
        i += 1
    doc.save(out_path)
    print("built", out_path)


if __name__ == '__main__':
    build(sys.argv[1], sys.argv[2])
